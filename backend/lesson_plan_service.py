import io
import json
import os
import re
from html import escape
from pathlib import Path
from typing import Any, Dict, Iterable, List, Tuple

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt
from docx.table import Table as DocxTable, _Cell
from docx.text.paragraph import Paragraph
from pypdf import PdfReader


LESSON_PLAN_OUTPUT_DIR = Path(__file__).parent / "generated" / "lesson-plans"
RULE_BASED_PROVIDER = "rule-based"
RULE_BASED_MODEL = "template-mapper-v1"
MAX_PDF_TEXT_CHARS = 24000
LABEL_ALIASES = {
    "lesson title": ["lesson", "title", "lesson title", "chapter", "topic"],
    "objectives": ["objectives", "lesson objectives", "aims", "learning objectives"],
    "aids": ["check list", "what you need for this lesson", "hardware", "software", "resources", "aids"],
    "lesson procedure": ["overview steps to follow", "steps to follow", "overview", "procedure", "lesson procedure"],
    "time pacing": ["duration", "time pacing", "periods", "timing"],
    "activities": ["activities", "project", "projects", "flashback", "steps to follow"],
    "group projects": ["project", "projects", "group projects"],
    "formative assessment": ["assessment", "formative assessment", "observation", "check", "evaluation"],
    "observation": ["observation", "assessment", "evaluation"],
    "warmup": ["warmup questions", "warm up questions", "warm up", "starter", "introduction"],
}


def sanitize_filename(filename: str, fallback: str) -> str:
    raw = (filename or "").strip()
    if not raw:
        raw = fallback
    raw = raw.replace("\\", "/").split("/")[-1]
    safe = re.sub(r"[^A-Za-z0-9._-]+", "-", raw).strip("-.")
    return safe or fallback


def ensure_docx_file(filename: str) -> str:
    safe = sanitize_filename(filename, "lesson-plan-template.docx")
    if not safe.lower().endswith(".docx"):
        raise ValueError("Please upload a .docx Word document.")
    return safe


def ensure_pdf_file(filename: str) -> str:
    safe = sanitize_filename(filename, "lesson-plan-source.pdf")
    if not safe.lower().endswith(".pdf"):
        raise ValueError("Please upload a PDF file.")
    return safe


def _normalize_text(value: Any) -> str:
    text = str(value or "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _truncate(text: str, limit: int) -> str:
    if len(text) <= limit:
        return text
    return f"{text[:limit].rstrip()}\n\n[truncated]"


def _normalize_key(value: Any) -> str:
    text = _normalize_text(value).lower()
    text = re.sub(r"[:：]", " ", text)
    text = re.sub(r"[^a-z0-9\u0600-\u06ff]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _looks_like_heading(text: str) -> bool:
    line = _normalize_text(text)
    if not line or len(line) > 80:
        return False
    if ":" in line or "：" in line:
        return False
    if re.search(r"[.!?]$", line):
        return False
    words = line.split()
    if len(words) > 8:
        return False
    return any(ch.isalpha() for ch in line)


def _looks_like_label(text: str) -> bool:
    line = _normalize_text(text)
    if not line or len(line) > 60:
        return False
    if ":" in line or "：" in line:
        return False
    if len(line.split()) > 6:
        return False
    return any(ch.isalpha() for ch in line)


def _extract_labeled_value(text: str) -> Tuple[str, str] | None:
    line = _normalize_text(text)
    if not line:
        return None
    match = re.match(r"^\s*([^:：]{1,60})\s*[:：]\s*(.+?)\s*$", line, flags=re.S)
    if not match:
        return None
    label = _normalize_text(match.group(1))
    value = _normalize_text(match.group(2))
    if not label or not value:
        return None
    return label, value


def _iter_block_items(parent: DocxDocument | _Cell) -> Iterable[Paragraph | DocxTable]:
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError("Unsupported parent for DOCX traversal")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield DocxTable(child, parent)


def _build_word_preview_html(document: DocxDocument, max_tables: int = 8) -> str:
    html_parts: List[str] = []
    table_count = 0
    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            text = _normalize_text(block.text)
            if text:
                html_parts.append(f"<p>{escape(text)}</p>")
            continue

        if table_count >= max_tables:
            continue
        table_count += 1
        rows_html: List[str] = []
        for row in block.rows:
            cells_html: List[str] = []
            for cell in row.cells:
                cell_text = "<br/>".join(
                    escape(_normalize_text(paragraph.text))
                    for paragraph in cell.paragraphs
                    if _normalize_text(paragraph.text)
                )
                cells_html.append(f"<td>{cell_text or '&nbsp;'}</td>")
            rows_html.append(f"<tr>{''.join(cells_html)}</tr>")
        html_parts.append(
            "<table class=\"lesson-plan-preview-table\"><tbody>"
            f"{''.join(rows_html)}"
            "</tbody></table>"
        )
    return "".join(html_parts) or "<p>No readable content found in this Word file.</p>"


def extract_docx_template(docx_bytes: bytes) -> Dict[str, Any]:
    document = Document(io.BytesIO(docx_bytes))
    editable_blocks: List[Dict[str, Any]] = []
    table_previews: List[Dict[str, Any]] = []
    paragraph_count = 0
    table_count = 0

    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            block_id = f"p-{paragraph_count}"
            text = _normalize_text(block.text)
            editable_blocks.append(
                {
                    "block_id": block_id,
                    "kind": "paragraph",
                    "text": text,
                    "style": (block.style.name if block.style is not None else "") or "",
                    "order": len(editable_blocks),
                    "context_label": "",
                }
            )
            paragraph_count += 1
            continue

        table_id = table_count
        table_count += 1
        row_text_cache: Dict[Tuple[int, int], str] = {}
        rows_preview: List[List[str]] = []
        for row_index, row in enumerate(block.rows):
            row_preview: List[str] = []
            for cell_index, cell in enumerate(row.cells):
                cell_texts: List[str] = []
                for para_index, paragraph in enumerate(cell.paragraphs):
                    block_id = f"t-{table_id}-r-{row_index}-c-{cell_index}-p-{para_index}"
                    text = _normalize_text(paragraph.text)
                    editable_blocks.append(
                        {
                            "block_id": block_id,
                            "kind": "table_cell",
                            "text": text,
                            "style": (paragraph.style.name if paragraph.style is not None else "") or "",
                            "order": len(editable_blocks),
                            "context_label": "",
                            "table_id": table_id,
                            "row_index": row_index,
                            "cell_index": cell_index,
                        }
                    )
                    if text:
                        cell_texts.append(text)
                cell_joined = "\n".join(cell_texts).strip()
                row_text_cache[(row_index, cell_index)] = cell_joined
                row_preview.append(cell_joined)
            rows_preview.append(row_preview)
        table_previews.append({"table_id": f"t-{table_id}", "rows": rows_preview})

        for item in editable_blocks:
            if item.get("kind") != "table_cell" or item.get("table_id") != table_id:
                continue
            row_index = item.get("row_index", 0)
            cell_index = item.get("cell_index", 0)
            if cell_index > 0:
                item["context_label"] = row_text_cache.get((row_index, cell_index - 1), "")

    non_empty_blocks = [block for block in editable_blocks if block["text"]]
    preview_text = "\n\n".join(block["text"] for block in non_empty_blocks[:20])
    return {
        "document": document,
        "editable_blocks": editable_blocks,
        "non_empty_blocks": non_empty_blocks,
        "paragraph_count": paragraph_count,
        "table_count": table_count,
        "preview_text": _truncate(preview_text, 4000),
        "preview_html": _build_word_preview_html(document),
        "table_previews": table_previews[:8],
    }


def extract_pdf_text(pdf_bytes: bytes) -> Dict[str, Any]:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    pages: List[str] = []
    for page in reader.pages:
        pages.append(_normalize_text(page.extract_text() or ""))
    joined = "\n\n".join(page for page in pages if page)
    lines = [_normalize_text(line) for line in joined.split("\n") if _normalize_text(line)]
    paragraphs = [_normalize_text(part) for part in re.split(r"\n\s*\n+", joined) if _normalize_text(part)]
    return {
        "page_count": len(reader.pages),
        "text": joined,
        "preview_text": _truncate(joined, 5000),
        "lines": lines,
        "paragraphs": paragraphs,
    }


def _extract_pdf_structured_content(pdf_text: str) -> Dict[str, Any]:
    trimmed_pdf = _truncate(pdf_text, MAX_PDF_TEXT_CHARS)
    lines = [_normalize_text(line) for line in trimmed_pdf.split("\n") if _normalize_text(line)]
    paragraphs = [_normalize_text(part) for part in re.split(r"\n\s*\n+", trimmed_pdf) if _normalize_text(part)]
    labeled_map: Dict[str, str] = {}
    ordered_units: List[str] = []
    seen_units: set[str] = set()

    def remember_unit(text: str) -> None:
        clean = _normalize_text(text)
        if clean and clean not in seen_units:
            ordered_units.append(clean)
            seen_units.add(clean)

    for index, line in enumerate(lines):
        labeled = _extract_labeled_value(line)
        if labeled:
            label, value = labeled
            labeled_map.setdefault(_normalize_key(label), value)
            remember_unit(value)
            continue
        if _looks_like_label(line) and index + 1 < len(lines):
            next_line = lines[index + 1]
            if next_line and len(next_line) > 10:
                labeled_map.setdefault(_normalize_key(line), next_line)
                remember_unit(next_line)

    for paragraph in paragraphs:
        labeled = _extract_labeled_value(paragraph) if "\n" not in paragraph else None
        if labeled:
            label, value = labeled
            labeled_map.setdefault(_normalize_key(label), value)
            remember_unit(value)
        else:
            remember_unit(paragraph)

    for index, line in enumerate(lines[:-1]):
        if _looks_like_heading(line) and not _extract_labeled_value(line):
            next_line = lines[index + 1]
            if next_line and len(next_line) > 10:
                labeled_map.setdefault(_normalize_key(line), next_line)

    return {
        "text": trimmed_pdf,
        "paragraphs": paragraphs,
        "lines": lines,
        "labeled_map": labeled_map,
        "ordered_units": ordered_units,
    }


def _find_section_index(lines: List[str], *phrases: str) -> int:
    normalized_lines = [_normalize_key(line) for line in lines]
    normalized_phrases = [_normalize_key(phrase) for phrase in phrases if _normalize_key(phrase)]
    if not normalized_phrases:
        return -1
    for index in range(len(normalized_lines)):
        windows = [
            normalized_lines[index],
            " ".join(normalized_lines[index:index + 2]).strip(),
            " ".join(normalized_lines[index:index + 3]).strip(),
        ]
        for phrase in normalized_phrases:
            if any(phrase and window.startswith(phrase) for window in windows):
                return index
    return -1


def _collect_list_items(lines: List[str]) -> List[str]:
    items: List[str] = []
    current = ""
    for raw_line in lines:
        line = _normalize_text(raw_line)
        if not line:
            continue
        normalized = _normalize_key(line)
        starts_new_item = bool(
            line.startswith(("-", "•"))
            or re.match(r"^(question \d+|for all students|for advanced students|teachers are asked to|students are asked to)\b", normalized)
            or re.match(r"^\d+ [a-z]", normalized)
            or "post activity" in normalized
        )
        if starts_new_item:
            if current:
                items.append(_normalize_text(current))
            current = line
        elif current:
            current = f"{current} {line.lstrip('-• ').strip()}".strip()
        else:
            current = line
    if current:
        items.append(_normalize_text(current))
    return [item for item in items if item]


def _fit_items_to_block_count(items: List[str], count: int) -> List[str]:
    normalized_items = [_normalize_text(item) for item in items if _normalize_text(item)]
    if count <= 0:
        return []
    if not normalized_items:
        return [""] * count
    if len(normalized_items) < count:
        return normalized_items + [""] * (count - len(normalized_items))
    if len(normalized_items) == count:
        return normalized_items
    head = normalized_items[:count - 1]
    tail = "\n".join(normalized_items[count - 1:])
    return head + [tail]


def _strip_question_prefix(text: str) -> str:
    normalized = _normalize_text(text)
    if not normalized:
        return ""
    return _normalize_text(re.sub(r"^\s*question\s*\d+\s*:\s*", "", normalized, flags=re.I))


def _merge_warmup_prompt(template_text: str, question_text: str) -> str:
    template = _normalize_text(template_text)
    question = _strip_question_prefix(question_text)
    if not template:
        return question
    if not question:
        return template
    match = re.match(r"^(.*?\bif we)\b.*$", template, flags=re.I)
    if match:
        prefix = match.group(1).rstrip()
        return f"{prefix} {question}".strip()
    return question


def _derive_header_topic(title: str) -> str:
    normalized = _normalize_text(title)
    if not normalized:
        return ""
    normalized = re.sub(r"\bLesson\s+\d+\b.*$", "", normalized, flags=re.I).strip(" -")
    normalized = re.sub(r"\bPages?\s+\d+\b.*$", "", normalized, flags=re.I).strip(" -")
    return _normalize_text(normalized)


def _extract_pdf_lesson_sections(lines: List[str]) -> Dict[str, Any]:
    if not lines:
        return {}

    def section_slice(start: int, end: int | None) -> List[str]:
        if start < 0:
            return []
        if end is None or end < 0:
            end = len(lines)
        return lines[start:end]

    lesson_objectives_idx = _find_section_index(lines, "lesson objectives")
    checklist_idx = _find_section_index(lines, "check list what you need", "what you need for this lesson")
    warmup_idx = _find_section_index(lines, "warmup questions", "warm up questions")
    overview_idx = _find_section_index(lines, "overview steps to follow", "steps to follow")
    hands_on_idx = _find_section_index(lines, "hands on")
    exercise_idx = _find_section_index(lines, "exercise project", "exercise project for all students")
    words_idx = _find_section_index(lines, "words to remember")
    challenge_idx = _find_section_index(lines, "challenge")

    title = _normalize_text(re.sub(r"\bPages?\s+\d+\b", "", lines[0], flags=re.I))
    header_topic = _derive_header_topic(title)

    objective_lines = section_slice(lesson_objectives_idx + 1, checklist_idx)
    duration_within_objectives_idx = _find_section_index(objective_lines, "duration")
    if duration_within_objectives_idx >= 0:
        objective_lines = objective_lines[:duration_within_objectives_idx]
    objective_intro = ""
    objective_bullets: List[str] = []
    if objective_lines:
        for index, line in enumerate(objective_lines):
            normalized = _normalize_key(line)
            if normalized.startswith("by the end of this lesson"):
                objective_intro = _normalize_text(line)
                objective_bullets = _collect_list_items(objective_lines[index + 1:])
                break
        if not objective_intro:
            objective_intro = _normalize_text(objective_lines[0])
            objective_bullets = _collect_list_items(objective_lines[1:])

    aids_lines = section_slice(checklist_idx, warmup_idx)
    prerequisite_idx = _find_section_index(aids_lines, "prerequisite")
    if prerequisite_idx >= 0:
        aids_lines = aids_lines[:prerequisite_idx]
    aids_content = [
        _normalize_text(line)
        for line in aids_lines
        if _normalize_text(line)
        and _normalize_key(line) not in {"check list what you need", "for this lesson"}
    ]
    aids = " ".join(aids_content)

    warmup_lines = section_slice(warmup_idx, overview_idx)
    warmup_questions = [
        item
        for item in _collect_list_items(warmup_lines)
        if not any(
            marker in _normalize_key(item)
            for marker in ("warmup", "warm up", "questions")
        )
    ]
    warmup_question = warmup_questions[0] if warmup_questions else ""

    procedure_lines = section_slice(overview_idx, hands_on_idx)
    overview_intro_lines: List[str] = []
    flashback_lines: List[str] = []
    project_lines: List[str] = []
    post_activity_lines: List[str] = []
    current_section = "overview"
    for line in procedure_lines:
        normalized = _normalize_key(line)
        if not normalized or normalized in {"overview", "steps to", "follow", "overview steps to follow"}:
            continue
        if normalized.startswith("1 flashback"):
            current_section = "flashback"
            continue
        if normalized.startswith("2 project"):
            current_section = "project"
            continue
        if "post activity" in normalized:
            current_section = "post_activity"
            continue
        if current_section == "overview":
            overview_intro_lines.append(line)
        elif current_section == "flashback":
            flashback_lines.append(line)
        elif current_section == "project":
            project_lines.append(line)
        else:
            post_activity_lines.append(line)

    hands_on_lines = section_slice(hands_on_idx, exercise_idx)
    teacher_lines: List[str] = []
    student_lines: List[str] = []
    current_group = ""
    for line in hands_on_lines:
        normalized = _normalize_key(line)
        if not normalized or normalized == "hands on":
            continue
        if normalized.startswith("teachers are asked to"):
            current_group = "teachers"
            continue
        if current_group == "teachers" and normalized == "show students":
            continue
        if normalized.startswith("students are asked to"):
            current_group = "students"
            continue
        if current_group == "teachers":
            teacher_lines.append(line)
        elif current_group == "students":
            student_lines.append(line)

    exercise_lines = section_slice(exercise_idx, words_idx)
    exercise_items = [
        item
        for item in _collect_list_items(exercise_lines)
        if "exercise project" not in _normalize_key(item)
    ]

    words_to_remember = ""
    if words_idx >= 0:
        words_line = _normalize_text(lines[words_idx])
        match = re.match(r"(?i)^Words to remember\s*(.+)$", words_line)
        if match:
            words_to_remember = _normalize_text(match.group(1))
        elif words_idx + 1 < len(lines):
            words_to_remember = _normalize_text(lines[words_idx + 1])

    challenge = ""
    if challenge_idx >= 0:
        challenge_line = _normalize_text(lines[challenge_idx])
        match = re.match(r"(?i)^Challenge\s*(.+)$", challenge_line)
        if match:
            challenge = _normalize_text(match.group(1))
        elif challenge_idx + 1 < len(lines):
            challenge = _normalize_text(lines[challenge_idx + 1])

    return {
        "title": title,
        "header_topic": header_topic,
        "objective_intro": objective_intro,
        "objective_bullets": objective_bullets,
        "aids": aids,
        "warmup_questions": warmup_questions,
        "warmup_question": warmup_question,
        "overview_intro": _normalize_text(" ".join(overview_intro_lines)),
        "flashback_items": _collect_list_items(flashback_lines),
        "project_items": _collect_list_items(project_lines),
        "post_activity_items": _collect_list_items(post_activity_lines),
        "teacher_items": _collect_list_items(teacher_lines),
        "student_items": _collect_list_items(student_lines),
        "exercise_items": exercise_items,
        "words_to_remember": words_to_remember,
        "challenge": challenge,
    }


def _build_targeted_section_replacements(
    template_blocks: List[Dict[str, Any]],
    pdf_sections: Dict[str, Any],
) -> Tuple[List[Dict[str, str]], set[str]]:
    rows: Dict[Tuple[int, int], Dict[int, List[Dict[str, Any]]]] = {}
    paragraph_blocks = [block for block in template_blocks if block.get("kind") == "paragraph"]
    for block in template_blocks:
        if block.get("kind") != "table_cell":
            continue
        key = (block.get("table_id", -1), block.get("row_index", -1))
        rows.setdefault(key, {}).setdefault(block.get("cell_index", 0), []).append(block)

    replacements: List[Dict[str, str]] = []
    handled: set[str] = set()

    def queue_blocks(blocks: List[Dict[str, Any]], values: List[str]) -> None:
        if not blocks:
            return
        fitted_values = _fit_items_to_block_count(values, len(blocks))
        for block, value in zip(blocks, fitted_values):
            current = _normalize_text(block.get("text"))
            next_value = _normalize_text(value)
            if next_value != current:
                replacements.append({"block_id": block["block_id"], "text": next_value})
            handled.add(block["block_id"])

    header_topic = pdf_sections.get("header_topic") or ""
    if header_topic:
        for block in paragraph_blocks:
            text = _normalize_text(block.get("text"))
            normalized = _normalize_key(text)
            if "week" in normalized and "strategy" in normalized:
                handled.add(block["block_id"])
                continue
            if "chapter" in normalized and "time pacing" in normalized:
                updated_text = re.sub(
                    r"(Stander:\s*[^ ]+\s+)(.*?)(\s+Time Pacing:.*)",
                    rf"\1{header_topic}\3",
                    text,
                    flags=re.I,
                )
                if updated_text != text:
                    replacements.append({"block_id": block["block_id"], "text": _normalize_text(updated_text)})
                handled.add(block["block_id"])
                break

    for (_table_id, _row_index), cells in sorted(rows.items()):
        left_text = "\n".join(_normalize_text(item.get("text")) for item in cells.get(0, []) if _normalize_text(item.get("text")))
        right_blocks = [item for item in cells.get(1, []) if _normalize_text(item.get("text"))]
        label = _normalize_key(left_text)
        if not label or not right_blocks:
            continue

        if "lesson title" in label:
            for block in right_blocks:
                handled.add(block["block_id"])
            continue

        if label == "objectives":
            objective_lines = pdf_sections.get("objective_bullets") or []
            if objective_lines:
                intro = "Students will know & be able to perform the following skills:"
                queue_blocks(right_blocks, [intro] + objective_lines)
                continue

        if label == "aids" and pdf_sections.get("aids"):
            queue_blocks(right_blocks, [pdf_sections["aids"]])
            continue

        if "lesson procedure" in label and "group projects" in label:
            warmup_heading = _normalize_text(right_blocks[0].get("text")) if right_blocks else ""
            warmup_template_line = _normalize_text(right_blocks[1].get("text")) if len(right_blocks) > 1 else ""
            warmup_body = _merge_warmup_prompt(
                warmup_template_line,
                pdf_sections.get("warmup_question") or "",
            )
            overview_intro = pdf_sections.get("overview_intro") or ""

            numbered_section_starts = [
                index
                for index, block in enumerate(right_blocks)
                if re.match(r"^[^\w]*\d+\s*[-.)]", _normalize_text(block.get("text")))
            ]
            if len(numbered_section_starts) >= 3:
                first_start, second_start, third_start = numbered_section_starts[:3]
                warmup_blocks = right_blocks[:first_start]
                flashback_blocks = right_blocks[first_start:second_start]
                project_blocks = right_blocks[second_start:third_start]
                post_activity_blocks = right_blocks[third_start:]

                queue_blocks(
                    warmup_blocks,
                    [
                        warmup_heading,
                        warmup_body,
                        overview_intro,
                    ],
                )
                queue_blocks(
                    flashback_blocks,
                    ["1. Flashback"] + (pdf_sections.get("flashback_items") or []),
                )
                queue_blocks(
                    project_blocks,
                    ["2. Project"] + (pdf_sections.get("project_items") or []),
                )
                post_items = pdf_sections.get("post_activity_items") or pdf_sections.get("teacher_items") or []
                queue_blocks(
                    post_activity_blocks,
                    ["Post-Activity"] + post_items,
                )
                continue

        if label == "coding ai":
            queue_blocks(right_blocks, [""] * len(right_blocks))
            continue

        if label == "assessment 5 min":
            if len(right_blocks) > 1:
                assessment_prompt = "\n".join((pdf_sections.get("exercise_items") or [])[:2]) or pdf_sections.get("challenge") or ""
                queue_blocks(
                    right_blocks,
                    [
                        _normalize_text(right_blocks[0].get("text")),
                        assessment_prompt,
                    ],
                )
            else:
                observation_text = "\n".join((pdf_sections.get("student_items") or [])[:3]) or "\n".join((pdf_sections.get("exercise_items") or [])[2:4])
                queue_blocks(right_blocks, [observation_text])
            continue

    return replacements, handled


def _best_pdf_match(label: str, labeled_map: Dict[str, str]) -> str | None:
    key = _normalize_key(label)
    if not key:
        return None
    candidate_keys = [key]
    for alias_key, aliases in LABEL_ALIASES.items():
        if alias_key in key or key in alias_key:
            candidate_keys.extend(_normalize_key(alias) for alias in aliases)
    expanded = []
    for candidate in candidate_keys:
        if candidate and candidate not in expanded:
            expanded.append(candidate)

    for candidate in expanded:
        if candidate in labeled_map:
            return labeled_map[candidate]

    best_value = None
    best_score = 0
    for candidate in expanded:
        label_tokens = set(candidate.split())
        if not label_tokens:
            continue
        for candidate_key, candidate_value in labeled_map.items():
            candidate_tokens = set(candidate_key.split())
            overlap = len(label_tokens & candidate_tokens)
            if not overlap:
                continue
            label_coverage = overlap / max(len(label_tokens), 1)
            candidate_coverage = overlap / max(len(candidate_tokens), 1)
            if label_coverage < 0.5 or candidate_coverage < 0.5:
                continue
            if overlap > best_score:
                best_score = overlap
                best_value = candidate_value
    return best_value if best_score >= 1 else None


def _replace_after_colon(template_text: str, new_value: str) -> str:
    match = re.match(r"^(\s*[^:：]{1,60}\s*[:：]\s*)(.+?)\s*$", template_text, flags=re.S)
    if not match:
        return new_value
    return f"{match.group(1)}{new_value}"


def _is_static_heading(block: Dict[str, Any]) -> bool:
    text = _normalize_text(block.get("text"))
    if not text:
        return True
    if block.get("kind") == "table_cell" and block.get("cell_index", 0) > 0 and block.get("context_label"):
        return False
    if _extract_labeled_value(text):
        return False
    if block.get("kind") == "table_cell" and block.get("cell_index", 0) == 0 and _looks_like_label(text):
        return True
    if _looks_like_heading(text):
        return True
    return len(text.split()) <= 3 and len(text) <= 32


def generate_lesson_plan_replacements(
    pdf_text: str,
    template_blocks: List[Dict[str, Any]],
    *,
    provider: str | None = None,
    model: str | None = None,
) -> Dict[str, Any]:
    _ = provider, model
    structured_pdf = _extract_pdf_structured_content(pdf_text)
    pdf_sections = _extract_pdf_lesson_sections(structured_pdf["lines"])
    labeled_map = structured_pdf["labeled_map"]
    ordered_units = structured_pdf["ordered_units"]
    queue_index = 0
    replacements, handled_block_ids = _build_targeted_section_replacements(template_blocks, pdf_sections)
    label_matches = 0
    sequential_matches = 0
    used_values: set[str] = set()

    for block in template_blocks:
        text = _normalize_text(block.get("text"))
        if not text or block.get("block_id") in handled_block_ids or _is_static_heading(block):
            continue

        replacement_text = None
        labeled = _extract_labeled_value(text)
        if labeled:
            label, _current_value = labeled
            matched = _best_pdf_match(label, labeled_map)
            if matched:
                replacement_text = _replace_after_colon(text, matched)
                label_matches += 1
                used_values.add(_normalize_text(matched))
        elif block.get("context_label"):
            matched = _best_pdf_match(block["context_label"], labeled_map)
            if matched:
                replacement_text = matched
                label_matches += 1
                used_values.add(_normalize_text(matched))
            else:
                candidate, queue_index = _next_ordered_unit(ordered_units, queue_index, used_values)
                if candidate:
                    replacement_text = candidate
                    sequential_matches += 1
                    used_values.add(candidate)
        else:
            candidate, queue_index = _next_ordered_unit(ordered_units, queue_index, used_values)
            if candidate:
                replacement_text = candidate
                sequential_matches += 1
                used_values.add(candidate)

        replacement_text = _normalize_text(replacement_text)
        if replacement_text and replacement_text != text:
            replacements.append({"block_id": block["block_id"], "text": replacement_text})

    summary = (
        f"Rule-based mapping updated {len(replacements)} blocks. "
        f"{label_matches} blocks were matched using detected PDF labels, and {sequential_matches} blocks were filled in template order."
    )
    return {
        "provider": RULE_BASED_PROVIDER,
        "model": RULE_BASED_MODEL,
        "summary": summary,
        "replacements": replacements,
    }


def _next_ordered_unit(ordered_units: List[str], start_index: int, used_values: set[str]) -> Tuple[str | None, int]:
    index = start_index
    while index < len(ordered_units):
        candidate = _normalize_text(ordered_units[index])
        index += 1
        if candidate and candidate not in used_values:
            return candidate, index
    return None, index


def _replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def _apply_compact_layout(paragraph: Paragraph, block: Dict[str, Any], text: str) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1.0

    row_index = block.get("row_index", -1)
    is_table_block = block.get("kind") == "table_cell"
    compact_rows = {3, 5, 6}
    size_pt = 9
    if is_table_block and row_index in compact_rows:
        size_pt = 8
    if len(_normalize_text(text)) > 180:
        size_pt = min(size_pt, 8)
    if len(_normalize_text(text)) > 320:
        size_pt = min(size_pt, 7)

    for run in paragraph.runs:
        run.font.size = Pt(size_pt)


def apply_replacements_to_document(docx_bytes: bytes, replacements: List[Dict[str, str]]) -> Tuple[bytes, int]:
    template_data = extract_docx_template(docx_bytes)
    document: DocxDocument = template_data["document"]
    block_lookup: Dict[str, Paragraph] = {}
    block_meta_lookup: Dict[str, Dict[str, Any]] = {
        block["block_id"]: block for block in template_data["editable_blocks"]
    }

    paragraph_count = 0
    table_count = 0
    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            block_lookup[f"p-{paragraph_count}"] = block
            paragraph_count += 1
            continue

        table_id = table_count
        table_count += 1
        for row_index, row in enumerate(block.rows):
            for cell_index, cell in enumerate(row.cells):
                for para_index, paragraph in enumerate(cell.paragraphs):
                    block_lookup[f"t-{table_id}-r-{row_index}-c-{cell_index}-p-{para_index}"] = paragraph

    updated = 0
    for replacement in replacements:
        block_id = replacement.get("block_id") or ""
        paragraph = block_lookup.get(block_id)
        if paragraph is None:
            continue
        replacement_text = replacement.get("text") or ""
        _replace_paragraph_text(paragraph, replacement_text)
        _apply_compact_layout(paragraph, block_meta_lookup.get(block_id, {}), replacement_text)
        updated += 1

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue(), updated


def save_generated_docx(user_id: str, filename: str, docx_bytes: bytes) -> Dict[str, str]:
    LESSON_PLAN_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document_id = sanitize_filename(user_id, "user").replace(".", "-") + "-" + sanitize_filename(os.urandom(6).hex(), "doc")
    safe_name = sanitize_filename(filename, "generated-lesson-plan.docx")
    output_path = LESSON_PLAN_OUTPUT_DIR / f"{document_id}.docx"
    output_path.write_bytes(docx_bytes)
    metadata_path = LESSON_PLAN_OUTPUT_DIR / f"{document_id}.json"
    metadata_path.write_text(json.dumps({"filename": safe_name}, ensure_ascii=False), encoding="utf-8")
    return {"document_id": document_id, "filename": safe_name, "path": str(output_path)}

